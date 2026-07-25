#!/usr/bin/env python3
"""
DOCX文档转换器
将DOCX文档转换为HTML格式
"""

import logging
from typing import Tuple, Optional

from .base import BaseConverter, HTMLWrapper, ConversionError

logger = logging.getLogger(__name__)


class DocxConverter(BaseConverter):
    """DOCX文档转换器"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = ['docx']
    
    def is_supported(self, file_type: str) -> bool:
        """检查文件类型是否支持转换"""
        return file_type.lower() in self.supported_formats
    
    def convert_to_html(self, file_path: str, file_type: str) -> Tuple[bool, Optional[str], Optional[str]]:
        """将DOCX文档转换为HTML"""
        try:
            # 验证文件
            is_valid, error_msg = self.validate_file(file_path)
            if not is_valid:
                return False, None, error_msg
            
            # 尝试使用mammoth库（如果安装了）
            try:
                import mammoth
                
                with open(file_path, "rb") as docx_file:
                    result = mammoth.convert_to_html(docx_file)
                    html_content = result.value
                    
                    # 包装成完整的HTML页面
                    full_html = self._wrap_docx_html(html_content, "Word文档预览")
                    return True, full_html, None
                    
            except ImportError:
                # 如果没有mammoth，使用基础的docx解析
                return self._convert_docx_basic(file_path)
                
        except Exception as e:
            logger.error(f"转换DOCX失败: {str(e)}")
            return False, None, str(e)
    
    def _convert_docx_basic(self, file_path: str) -> Tuple[bool, Optional[str], Optional[str]]:
        """基础DOCX转换（提取文本内容）"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            content_parts = []
            
            # 处理段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(f"<p>{HTMLWrapper.escape_html(paragraph.text)}</p>")
            
            # 处理表格
            for table in doc.tables:
                content_parts.append("<table border='1' style='border-collapse: collapse; width: 100%;'>")
                for row in table.rows:
                    content_parts.append("<tr>")
                    for cell in row.cells:
                        content_parts.append(f"<td style='padding: 8px; border: 1px solid #ddd;'>{HTMLWrapper.escape_html(cell.text)}</td>")
                    content_parts.append("</tr>")
                content_parts.append("</table><br>")
            
            html_content = "\n".join(content_parts)
            full_html = self._wrap_docx_html(html_content, "Word文档预览")
            
            return True, full_html, None
            
        except Exception as e:
            logger.error(f"基础DOCX转换失败: {str(e)}")
            return False, None, str(e)
    
    def _wrap_docx_html(self, content: str, title: str) -> str:
        """为DOCX内容包装HTML页面，优化移动端和微信浏览器兼容性"""
        escaped_title = HTMLWrapper.escape_html(title)
        
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0, user-scalable=yes, minimum-scale=0.5, maximum-scale=3.0">
            <meta name="format-detection" content="telephone=no">
            <meta name="apple-mobile-web-app-capable" content="yes">
            <meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
            <title>{escaped_title}</title>
            <style>
                * {{
                    box-sizing: border-box;
                }}
                body {{
                    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
                    line-height: 1.5;
                    max-width: 1200px;
                    margin: 0 auto;
                    padding: 5px;
                    background: white;
                    color: #333;
                    font-size: 14px;
                    -webkit-text-size-adjust: 100%;
                    -webkit-font-smoothing: antialiased;
                }}
                .container {{
                    background: white;
                    padding: 8px;
                    margin: 0;
                }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 10px 0;
                    font-size: 13px;
                    overflow-x: auto;
                    display: block;
                    white-space: nowrap;
                }}
                th, td {{
                    padding: 6px 8px;
                    border: 1px solid #ddd;
                    text-align: left;
                    vertical-align: top;
                    word-wrap: break-word;
                    max-width: 200px;
                    font-size: 13px;
                }}
                th {{
                    background-color: #f8f9fa;
                    font-weight: 600;
                    color: #495057;
                }}
                tr:nth-child(even) {{
                    background-color: #f8f9fa;
                }}
                h1, h2, h3, h4, h5, h6 {{
                    color: #2c3e50;
                    margin: 15px 0 8px 0;
                    word-wrap: break-word;
                }}
                h1 {{ font-size: 18px; }}
                h2 {{ font-size: 16px; }}
                h3 {{ font-size: 15px; }}
                h4 {{ font-size: 14px; }}
                h5 {{ font-size: 13px; }}
                h6 {{ font-size: 12px; }}
                p {{
                    margin: 8px 0;
                    word-wrap: break-word;
                    line-height: 1.5;
                    font-size: 14px;
                }}
                /* 移动端优化 */
                @media (max-width: 768px) {{
                    body {{
                        padding: 3px;
                        font-size: 13px;
                    }}
                    .container {{
                        padding: 6px;
                    }}
                    table {{
                        font-size: 12px;
                        display: table;
                        white-space: normal;
                    }}
                    th, td {{
                        padding: 5px 6px;
                        max-width: 120px;
                        word-break: break-all;
                        font-size: 12px;
                    }}
                    h1 {{ font-size: 16px; }}
                    h2 {{ font-size: 15px; }}
                    h3 {{ font-size: 14px; }}
                    h4 {{ font-size: 13px; }}
                    p {{ font-size: 13px; }}
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="content">
                    {content}
                </div>
            </div>

            <script>
                // 优化表格显示
                document.addEventListener('DOMContentLoaded', function() {{
                    const tables = document.querySelectorAll('table');
                    tables.forEach(function(table) {{
                        // 为表格添加滚动容器
                        const wrapper = document.createElement('div');
                        wrapper.style.overflowX = 'auto';
                        wrapper.style.marginBottom = '15px';
                        table.parentNode.insertBefore(wrapper, table);
                        wrapper.appendChild(table);
                    }});
                }});
            </script>
        </body>
        </html>
        """
